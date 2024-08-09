""" 
Module Docstring not implemented yet
"""

import os
import shutil
import random
import subprocess
from datetime import date

from docx import Document
from docx.shared import Inches
from docx2pdf import convert

import logger as log
from Config import Config

class Invoice(Config):
    """ 
    Class Docstring not implemented yet
    """
    def __init__(self, logger) -> None:
        super().__init__(logger=logger)

    def write_image(self, paypal_email: str, total: float =None,
                    hourly_rate: float =20.00, hours_worked: float =0.00):
        """ 
        Method Dostring not implemented yet
        """
        shutil.copy("./Invoice_Template.docx", "./invoice_final.docx")

        doc = Document("./invoice_final.docx")

        page_width = doc.sections[0].page_width.inches
        left_margin = doc.sections[0].left_margin.inches
        right_margin = doc.sections[0].right_margin.inches
        max_width = page_width - left_margin - right_margin

        for picture in os.listdir(f"{self.folder_src()}/temp/"):
            if str(picture).startswith("table"):
                doc.add_picture(f"{self.folder_src()}/temp/{picture}", width=Inches(max_width))

        personal_cost = hourly_rate * hours_worked

        return self.write_text(doc, total, personal_cost, paypal_email)

    def write_text(self, doc, total, cost, paypal_email):
        """ 
        Method Dostring not implemented yet
        """
        doc.add_paragraph()

        p1 = doc.add_paragraph()
        p1.paragraph_format.tab_stops.add_tab_stop(Inches(1)) 
        p1.add_run("\tKosten Teile:").bold = True
        p1.add_run(f" {total} €")

        p2 = doc.add_paragraph()
        p2.paragraph_format.tab_stops.add_tab_stop(Inches(1))
        p2.add_run("\tKosten Lohn:").bold = True
        p2.add_run(f" {cost} €")

        p3 = doc.add_paragraph()
        p3.paragraph_format.tab_stops.add_tab_stop(Inches(1))
        p3.add_run("\tGesamtkosten:").bold = True
        p3.add_run(f" {round(total+cost, 2)} €")

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Book Antiqua"
                run.font.size = Inches(0.2)

        self.__fill_variables(doc=doc, paypal=paypal_email)
        doc.save("./invoice_final.docx")

        try:
            convert("./invoice_final.docx", "./invoice_final.pdf")
        except NotImplementedError:
            subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", "./invoice_final.docx", "--outdir", "./"], check=True)


    def __fill_variables(self, doc, paypal: str = None):
        data = {
            "[Date]": date.today().strftime("%d-%m-%Y"),
            "[Invoice-Number]": date.today().strftime("%d-%m-%Y"),
            "[PayPal]": paypal
        }

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)


if __name__ == "__main__":
    invoice = Invoice(logger=log.get_logger())
